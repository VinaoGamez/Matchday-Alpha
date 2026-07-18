"""Convert a Markdown doc to .docx (headings, lists, tables, code)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, size=11, bold=False, italic=False, code=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    font_name = "Consolas" if code else "Calibri"
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    if color:
        run.font.color.rgb = color


def add_inline(paragraph, content, base_size=11):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for part in pattern.split(content):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=base_size - 1, code=True, color=RGBColor(0x1A, 0x3A, 0x4A))
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=base_size)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=9, code=True, color=RGBColor(0x20, 0x40, 0x50))
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F3F6F8")
    shade.set(qn("w:val"), "clear")
    paragraph._p.get_or_add_pPr().append(shade)


def parse_table(doc, table_buf):
    rows = []
    for raw in table_buf:
        if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", raw.strip()):
            continue
        cells = [cell.strip() for cell in raw.strip().strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.rows[i].cells[j]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            value = row[j] if j < len(row) else ""
            add_inline(paragraph, value, base_size=10)
            if i == 0:
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph()


def convert(src: Path, out: Path) -> None:
    text = src.read_text(encoding="utf-8")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[str] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            parse_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if "|" in line and line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue

        flush_table()

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.strip() == "---":
            paragraph = doc.add_paragraph("─" * 40)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=8, color=RGBColor(0x99, 0x99, 0x99))
        elif re.match(r"^>\s?", line):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            add_inline(paragraph, line.lstrip("> ").strip())
            for run in paragraph.runs:
                run.italic = True
        elif re.match(r"^[-*]\s+", line):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, re.sub(r"^[-*]\s+", "", line))
        elif re.match(r"^\d+\.\s+", line):
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", line))
        else:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, line.strip())
        i += 1

    flush_table()
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out.resolve()} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    source = Path(sys.argv[1] if len(sys.argv) > 1 else "docs/08-JOGADAS-E-RITMO.md")
    target = Path(sys.argv[2] if len(sys.argv) > 2 else "docs/08-JOGADAS-E-RITMO.docx")
    convert(source, target)
